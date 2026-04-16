from __future__ import annotations

import base64
import os
import sys
import tempfile
import unittest
from pathlib import Path


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from local_meeting_ai_runtime.artifact_exporter import MeetingArtifactExporter


class ArtifactExporterRendererProfileTest(unittest.TestCase):
    def test_export_summary_bundle_delegates_to_html_renderer_when_selected(self) -> None:
        previous = os.environ.get("DELEGATE_MEETING_ARTIFACT_PDF_RENDERER")
        os.environ["DELEGATE_MEETING_ARTIFACT_PDF_RENDERER"] = "html"
        try:
            exporter = MeetingArtifactExporter()
            with tempfile.TemporaryDirectory() as temp_dir:
                summary_md = Path(temp_dir) / "summary.md"
                summary_md.write_text("# Summary\n", encoding="utf-8")

                captured: dict[str, object] = {}

                def fake_render(
                    summary_markdown_path: Path,
                    *,
                    briefing: dict[str, object],
                    rendering_policy: dict[str, object] | None = None,
                    postprocess_requests: list[dict[str, object]] | None = None,
                ) -> list[dict[str, str]]:
                    captured["path"] = summary_markdown_path
                    captured["briefing"] = briefing
                    captured["rendering_policy"] = rendering_policy or {}
                    captured["postprocess_requests"] = postprocess_requests or []
                    return [{"format": "html", "path": str(summary_markdown_path.with_suffix(".html"))}]

                exporter._html_renderer.render_summary_bundle = fake_render  # type: ignore[method-assign]
                exports = exporter.export_summary_bundle(
                    summary_md,
                    briefing={"title": "HTML Path"},
                    rendering_policy={"renderer_theme_name": "NAVER"},
                    postprocess_requests=[{"title": "Visual"}],
                )

            self.assertEqual(exports, [{"format": "html", "path": str(summary_md.with_suffix(".html"))}])
            self.assertEqual(captured["path"], summary_md)
            self.assertEqual(captured["briefing"], {"title": "HTML Path"})
            self.assertEqual(captured["rendering_policy"], {"renderer_theme_name": "NAVER"})
            self.assertEqual(captured["postprocess_requests"], [{"title": "Visual"}])
        finally:
            if previous is None:
                os.environ.pop("DELEGATE_MEETING_ARTIFACT_PDF_RENDERER", None)
            else:
                os.environ["DELEGATE_MEETING_ARTIFACT_PDF_RENDERER"] = previous

    def test_renderer_profile_settings_supports_brand_like_color_overrides(self) -> None:
        exporter = MeetingArtifactExporter()

        profile = exporter._renderer_profile_settings(
            "formal",
            rendering_policy={
                "renderer_theme_name": "kakao-like",
                "renderer_primary_color": "#FEE500",
                "renderer_accent_color": "3C1E1E",
                "renderer_neutral_color": "4A4A4A",
            },
        )

        self.assertEqual(profile["heading1_color"], "FEE500")
        self.assertEqual(profile["heading2_color"], "FEE500")
        self.assertEqual(profile["heading3_color"], "3C1E1E")
        self.assertEqual(profile["body_color"], "4A4A4A")
        self.assertEqual(profile["table_header_fill"], "FEE500")
        self.assertEqual(profile["table_label_fill"], "FFFDEB")

    def test_renderer_profile_settings_keeps_builtin_palette_without_overrides(self) -> None:
        exporter = MeetingArtifactExporter()

        profile = exporter._renderer_profile_settings("report", rendering_policy={})

        self.assertEqual(profile["heading1_color"], "1E4E79")
        self.assertEqual(profile["heading2_color"], "245D91")
        self.assertEqual(profile["heading3_color"], "2F6FA8")
        self.assertEqual(profile["table_header_fill"], "245D91")

    def test_resolve_postprocess_image_assets_accepts_existing_local_images(self) -> None:
        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / "visual.png"
            image_path.write_bytes(b"not-a-real-image-but-a-real-path")

            assets = exporter._resolve_postprocess_image_assets(
                [
                    {
                        "kind": "image_brief",
                        "title": "카카오 느낌 참고 이미지",
                        "instruction": "브랜드 무드 참고용 시각 자료",
                        "caption": "카카오 느낌 무드 보드",
                        "placement_notes": "핵심 논의 주제 섹션 안에 배치",
                        "target_heading": "핵심 논의 주제",
                        "image_path": str(image_path),
                    }
                ],
                base_dir=Path(temp_dir),
            )

        self.assertEqual(len(assets), 1)
        self.assertEqual(assets[0]["title"], "카카오 느낌 참고 이미지")
        self.assertEqual(assets[0]["caption"], "카카오 느낌 무드 보드")
        self.assertEqual(assets[0]["placement_notes"], "핵심 논의 주제 섹션 안에 배치")
        self.assertEqual(assets[0]["target_heading"], "핵심 논의 주제")
        self.assertTrue(assets[0]["path"].endswith("visual.png"))

    def test_resolve_postprocess_image_assets_accepts_session_relative_visual_paths(self) -> None:
        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            session_dir = Path(temp_dir) / "session-1"
            visuals_dir = session_dir / "visuals"
            visuals_dir.mkdir(parents=True, exist_ok=True)
            image_path = visuals_dir / "kakao-1.png"
            image_path.write_bytes(b"fake-image")

            assets = exporter._resolve_postprocess_image_assets(
                [
                    {
                        "kind": "image_brief",
                        "title": "카카오 무드 참고 이미지",
                        "instruction": "카카오 분위기의 보조 시각 자료",
                        "image_path": "visuals/kakao-1.png",
                    }
                ],
                base_dir=session_dir,
            )

        self.assertEqual(len(assets), 1)
        self.assertTrue(assets[0]["path"].endswith("visuals\\kakao-1.png") or assets[0]["path"].endswith("visuals/kakao-1.png"))

    def test_renderer_profile_settings_supports_font_overrides(self) -> None:
        exporter = MeetingArtifactExporter()

        profile = exporter._renderer_profile_settings(
            "formal",
            rendering_policy={
                "renderer_title_font": "Malgun Gothic",
                "renderer_heading_font": "Malgun Gothic",
                "renderer_body_font": "Malgun Gothic",
            },
        )

        self.assertEqual(profile["title_font"], "Malgun Gothic")
        self.assertEqual(profile["heading_font"], "Malgun Gothic")
        self.assertEqual(profile["body_font"], "Malgun Gothic")

    def test_renderer_design_settings_supports_direct_cover_controls(self) -> None:
        exporter = MeetingArtifactExporter()
        profile = exporter._renderer_profile_settings(
            "formal",
            rendering_policy={
                "renderer_primary_color": "#FEE500",
            },
        )

        design = exporter._renderer_design_settings(
            rendering_policy={
                "renderer_cover_align": "center",
                "renderer_surface_tint_color": "#FFF8CC",
                "renderer_cover_kicker": "COLLABORATION PROPOSAL",
                "renderer_primary_color": "#FEE500",
            },
            profile=profile,
        )

        self.assertEqual(design["cover_align"], "center")
        self.assertEqual(design["section_band_fill"], "")
        self.assertEqual(design["cover_fill"], "FFF8CC")
        self.assertEqual(design["cover_kicker"], "COLLABORATION PROPOSAL")

    def test_renderer_design_settings_supports_direct_surface_overrides_without_layout_menu(self) -> None:
        exporter = MeetingArtifactExporter()
        profile = exporter._renderer_profile_settings("formal", rendering_policy={})

        design = exporter._renderer_design_settings(
            rendering_policy={
                "renderer_cover_fill_color": "F7F1D0",
                "renderer_section_panel_fill_color": "F9F5E8",
                "renderer_overview_panel_fill_color": "F4EFE1",
            },
            profile=profile,
        )

        self.assertEqual(design["cover_fill"], "F7F1D0")
        self.assertEqual(design["section_panel_fill"], "F9F5E8")
        self.assertEqual(design["overview_panel_fill"], "F4EFE1")

    def test_polish_summary_docx_applies_soft_panel_structures(self) -> None:
        from docx import Document

        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "soft-summary.docx"
            document = Document()
            document.add_paragraph("Soft Design Title", style="Heading 1")
            document.add_paragraph("회의 개요", style="Heading 2")
            document.add_paragraph("회의 일시: 2026-04-11 10:00 KST")
            document.add_paragraph("참석자: Alice, Bob")
            document.add_paragraph("핵심 논의 주제", style="Heading 2")
            document.add_paragraph("1. 첫 번째 논의", style="Heading 3")
            document.add_paragraph("세부 내용입니다.")
            document.save(docx_path)

            exporter._polish_summary_docx(
                docx_path,
                renderer_profile="formal",
                rendering_policy={
                    "renderer_surface_tint_color": "FFF8CC",
                    "renderer_cover_kicker": "COLLABORATION PROPOSAL",
                    "renderer_section_band_fill_color": "FFF2B3",
                    "renderer_section_panel_fill_color": "FFF8E1",
                    "renderer_overview_panel_fill_color": "FFF8E1",
                    "overview_heading": "회의 개요",
                    "overview_datetime_label": "회의 일시",
                    "overview_participants_label": "참석자",
                },
            )

            polished = Document(docx_path)

        table_texts = ["\n".join(cell.text for row in table.rows for cell in row.cells) for table in polished.tables]
        self.assertTrue(any("COLLABORATION PROPOSAL" in text for text in table_texts))
        self.assertTrue(any("회의 일시" in text and "참석자" in text for text in table_texts))
        self.assertTrue(any("첫 번째 논의" in text and "세부 내용입니다." in text for text in table_texts))
        self.assertTrue(any(len(table.columns) == 1 for table in polished.tables))

    def test_polish_summary_docx_applies_structured_table_and_accent_bar(self) -> None:
        from docx import Document

        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "structured-summary.docx"
            document = Document()
            document.add_paragraph("Structured Design Title", style="Heading 1")
            document.add_paragraph("회의 개요", style="Heading 2")
            document.add_paragraph("회의 일시: 2026-04-11 10:00 KST")
            document.add_paragraph("참석자: Alice, Bob")
            document.add_paragraph("핵심 논의 주제", style="Heading 2")
            document.add_paragraph("1. 첫 번째 논의", style="Heading 3")
            document.add_paragraph("세부 내용입니다.")
            document.save(docx_path)

            exporter._polish_summary_docx(
                docx_path,
                renderer_profile="formal",
                rendering_policy={
                    "renderer_section_border_color": "245D91",
                    "renderer_section_accent_fill_color": "245D91",
                    "renderer_overview_label_fill_color": "E8EEF5",
                    "renderer_overview_value_fill_color": "FFFFFF",
                    "overview_heading": "회의 개요",
                    "overview_datetime_label": "회의 일시",
                    "overview_participants_label": "참석자",
                },
            )

            polished = Document(docx_path)

        two_column_tables = [table for table in polished.tables if len(table.columns) == 2]
        self.assertTrue(two_column_tables)
        self.assertTrue(any(table.cell(0, 0).text == "회의 일시" for table in two_column_tables))
        self.assertTrue(any("첫 번째 논의" in table.cell(0, 1).text for table in two_column_tables))

    def test_polish_summary_docx_places_inline_visuals_inside_target_section(self) -> None:
        from docx import Document

        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "inline-visual.docx"
            image_path = Path(temp_dir) / "inline-visual.png"
            image_path.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0ioAAAAASUVORK5CYII="
                )
            )

            document = Document()
            document.add_paragraph("Inline Visual Title", style="Heading 1")
            document.add_paragraph("?듭떖 ?쇱쓽 二쇱젣", style="Heading 2")
            document.add_paragraph("Section One", style="Heading 3")
            document.add_paragraph("This section should receive the generated visual inline.")
            document.save(docx_path)

            exporter._polish_summary_docx(
                docx_path,
                renderer_profile="formal",
                rendering_policy={
                    "postprocess_requests_heading": "Visual Appendix",
                },
                postprocess_requests=[
                    {
                        "kind": "image_brief",
                        "title": "Section Visual",
                        "instruction": "Inline supporting visual",
                        "caption": "Supports Section One",
                        "placement_notes": "Section One 섹션 안쪽 배치",
                        "target_heading": "Section One",
                        "image_path": str(image_path),
                    }
                ],
            )

            polished = Document(docx_path)

        self.assertEqual(len(polished.inline_shapes), 1)
        self.assertFalse(any(paragraph.text == "Visual Appendix" for paragraph in polished.paragraphs))
        table_texts = ["\n".join(cell.text for row in table.rows for cell in row.cells) for table in polished.tables]
        self.assertTrue(any("Section Visual" in text and "Supports Section One" in text for text in table_texts))

    def test_polish_summary_docx_places_overview_visual_before_next_block(self) -> None:
        from docx import Document

        exporter = MeetingArtifactExporter()
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "overview-visual.docx"
            image_path = Path(temp_dir) / "overview-visual.png"
            image_path.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0ioAAAAASUVORK5CYII="
                )
            )

            document = Document()
            document.add_paragraph("Overview Visual Title", style="Heading 1")
            document.add_paragraph("Overview", style="Heading 2")
            document.add_paragraph("Date: 2026-04-11 10:00 KST")
            document.add_paragraph("Participants: Alice, Bob")
            document.add_paragraph("Executive Summary", style="Heading 2")
            document.add_paragraph("Summary body")
            document.save(docx_path)

            exporter._polish_summary_docx(
                docx_path,
                renderer_profile="formal",
                rendering_policy={
                    "overview_heading": "Overview",
                    "executive_summary_heading": "Executive Summary",
                },
                postprocess_requests=[
                    {
                        "kind": "image_brief",
                        "title": "Overview Visual",
                        "instruction": "Overview supporting visual",
                        "caption": "Supports overview",
                        "placement_notes": "Overview 섹션 안쪽 배치",
                        "target_heading": "Overview",
                        "image_path": str(image_path),
                    }
                ],
            )

            polished = Document(docx_path)

        self.assertEqual(len(polished.inline_shapes), 1)
        table_texts = ["\n".join(cell.text for row in table.rows for cell in row.cells) for table in polished.tables]
        self.assertTrue(any("Overview Visual" in text and "Supports overview" in text for text in table_texts))
        self.assertFalse(any(paragraph.text == "후속 시각 자료" for paragraph in polished.paragraphs))

    def test_postprocess_image_width_inches_is_policy_controlled_and_clamped(self) -> None:
        exporter = MeetingArtifactExporter()

        self.assertEqual(exporter._postprocess_image_width_inches({"postprocess_image_width_inches": "6.4"}), 6.4)
        self.assertEqual(exporter._postprocess_image_width_inches({"postprocess_image_width_inches": "99"}), 7.2)
        self.assertEqual(exporter._postprocess_image_width_inches({"postprocess_image_width_inches": "0.5"}), 1.5)
        self.assertEqual(exporter._postprocess_image_width_inches({"postprocess_image_width_inches": "wide"}), 5.9)


if __name__ == "__main__":
    unittest.main()
